import os
from typing import Callable, Dict

import re

def clean_text_spacing(text: str) -> str:
    # Collapse single newlines to spaces, but PRESERVE them when the following
    # line starts with a list marker (-, *, •, or digit + . / )) or a page
    # marker (<<<PAGE:N>>>) so structural information survives extraction.
    _LIST_START = r'[ \t]*(?:[-*•]|\d+[.)]) '
    _PAGE_MARKER = r'<<<PAGE:\d+>>>'
    text = re.sub(
        rf'(?<!\n)\n(?!\n)(?!{_LIST_START})(?!{_PAGE_MARKER})',
        ' ',
        text,
    )
    # Collapse multiple spaces / tabs
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse 3+ newlines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# Lazy imports inside helpers so you don't pay the cost unless needed.

HEADER_FOOTER_PCT = 0.05  # top/bottom band to drop on PDFs


def _read_txt(path: str, encoding: str = "utf-8") -> str:
    with open(path, "r", encoding=encoding, errors="replace") as f:
        return f.read()


_PAGE_MARKER_PREFIX = "<<<PAGE:"
_PAGE_MARKER_SUFFIX = ">>>"


def _make_page_marker(page_num: int) -> str:
    return f"{_PAGE_MARKER_PREFIX}{page_num}{_PAGE_MARKER_SUFFIX}"


def _read_pdf_pymupdf(path: str, top_pct: float = HEADER_FOOTER_PCT, bottom_pct: float = HEADER_FOOTER_PCT) -> str:
    """Extract text blocks from a PDF using PyMuPDF, removing headers/footers by position.

    Each page's content is prefixed with a <<<PAGE:N>>> marker so the chunker can
    track which source page each chunk came from.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    try:
        pages_out = []
        for page_num, page in enumerate(doc, start=1):
            height = page.rect.height
            top_y = height * top_pct
            bottom_y = height * (1 - bottom_pct)

            # Each block: (x0, y0, x1, y1, text, block_no, block_type, ...)
            blocks_out = []
            for x0, y0, x1, y1, text, *_ in page.get_text("blocks"):
                if y1 < top_y or y0 > bottom_y:
                    continue
                text = (text or "").strip()
                if text:
                    blocks_out.append(text)

            if blocks_out:
                pages_out.append(
                    _make_page_marker(page_num) + "\n\n" + "\n\n".join(blocks_out)
                )

        return "\n\n".join(pages_out).strip()
    finally:
        doc.close()


def _read_pdf_pypdf2(path: str) -> str:
    """Fallback PDF reader when PyMuPDF yields nothing or isn't suitable."""
    import PyPDF2

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return "\n".join((p.extract_text() or "").strip() for p in reader.pages).strip()


def _read_docx(path: str) -> str:
    """Extract text from a DOCX file including tables.

    Detects explicit page breaks (w:br type="page") to emit <<<PAGE:N>>> markers
    so the chunker can tag each chunk with its source page number.
    """
    import docx  # python-docx
    from docx.oxml.ns import qn

    doc = docx.Document(path)
    parts: list[str] = []
    current_page = 1

    # Start with a page marker for page 1
    parts.append(_make_page_marker(current_page))

    # --- Body paragraphs ---
    for p in doc.paragraphs:
        # Detect explicit page breaks inside paragraph runs
        for run in p.runs:
            for br in run._r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    current_page += 1
                    parts.append(_make_page_marker(current_page))

        if p.text.strip():
            parts.append(p.text.strip())

    # --- Tables → pipe-delimited rows ---
    # python-docx repeats the value of merged cells for every position they occupy;
    # deduplicate adjacent identical values to avoid noise.
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            deduped: list[str] = [cells[0]] if cells else []
            for cell in cells[1:]:
                if cell != deduped[-1]:
                    deduped.append(cell)
            row_text = " | ".join(deduped)
            if row_text.strip(" |"):
                table_lines.append(row_text)
        if table_lines:
            parts.append("\n".join(table_lines))

    return "\n\n".join(parts).strip()


def _parse_file(file_path: str) -> str:
    """
    Parse a file and return its main textual content.

    Supported:
      - .txt     (UTF-8, with 'replace' for bad bytes)
      - .pdf     (PyMuPDF for layout-aware extraction + header/footer removal; PyPDF2 fallback)
      - .doc/.docx (python-docx)
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    readers: Dict[str, Callable[[str], str]] = {
        ".txt": _read_txt,
        ".doc": _read_docx,
        ".docx": _read_docx,
    }

    if ext == ".pdf":
        try:
            text = _read_pdf_pymupdf(file_path)
            if text:
                return text
        except Exception:
            # Fall through to PyPDF2 if PyMuPDF fails
            pass
        # PyMuPDF was empty or failed → try PyPDF2
        return _read_pdf_pypdf2(file_path)

    if ext in readers:
        return readers[ext](file_path)

    raise ValueError(f"Unsupported file type: {ext}")

def parse_file(file_path: str) -> str:
    """Wrapper to parse a file and clean its text spacing."""
    text = _parse_file(file_path)
    return clean_text_spacing(text)